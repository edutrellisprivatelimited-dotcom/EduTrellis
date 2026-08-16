"""Extracts plain text from an uploaded PDF/DOCX/CSV/TXT so it can be fed to
the AI chat as normal text — no OCR, no external service, just the
document's own text layer. Scanned/image-only PDFs have no text layer to
extract and will raise ExtractError rather than return nothing useful.
"""
import csv
import io
import os

from docx import Document as DocxDocument
from pypdf import PdfReader

MAX_CHARS = 15_000     # extracted text is capped before it ever reaches the model
MAX_PDF_PAGES = 50
MAX_CSV_ROWS = 500


class ExtractError(Exception):
    """Raised for any extraction failure with a message safe to show the user."""


def extract_pdf(file_bytes):
    try:
        reader = PdfReader(io.BytesIO(file_bytes))
    except Exception as e:
        raise ExtractError(f"Could not read that PDF: {e}")

    if reader.is_encrypted:
        try:
            reader.decrypt('')
        except Exception:
            pass
        if reader.is_encrypted:
            raise ExtractError('That PDF is password-protected — remove the password and try again.')

    pages = []
    for page in reader.pages[:MAX_PDF_PAGES]:
        try:
            pages.append(page.extract_text() or '')
        except Exception:
            continue
    text = '\n\n'.join(p for p in pages if p.strip())
    if not text.strip():
        raise ExtractError(
            "Couldn't find any readable text in that PDF — it may be a scanned/image-only document, "
            "which isn't supported yet."
        )
    return text


def extract_docx(file_bytes):
    try:
        doc = DocxDocument(io.BytesIO(file_bytes))
    except Exception as e:
        raise ExtractError(f"Could not read that document: {e}")

    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            parts.append(' | '.join(cell.text for cell in row.cells))
    text = '\n'.join(parts)
    if not text.strip():
        raise ExtractError('That document appears to be empty.')
    return text


def extract_csv(file_bytes):
    try:
        text_data = file_bytes.decode('utf-8-sig', errors='replace')
        rows = list(csv.reader(io.StringIO(text_data)))
    except Exception as e:
        raise ExtractError(f"Could not read that CSV: {e}")
    if not rows:
        raise ExtractError('That CSV appears to be empty.')
    truncated_rows = rows[:MAX_CSV_ROWS]
    lines = [', '.join(cell.strip() for cell in row) for row in truncated_rows]
    if len(rows) > MAX_CSV_ROWS:
        lines.append(f'... ({len(rows) - MAX_CSV_ROWS} more rows not shown)')
    return '\n'.join(lines)


def extract_txt(file_bytes):
    try:
        text = file_bytes.decode('utf-8-sig', errors='replace')
    except Exception as e:
        raise ExtractError(f"Could not read that file: {e}")
    if not text.strip():
        raise ExtractError('That file appears to be empty.')
    return text


_EXTRACTORS = {
    '.pdf': extract_pdf,
    '.docx': extract_docx,
    '.csv': extract_csv,
    '.txt': extract_txt,
}

SUPPORTED_EXTENSIONS = tuple(_EXTRACTORS.keys())


def extract(filename, file_bytes):
    """Returns (text, truncated). Raises ExtractError on any failure."""
    ext = os.path.splitext(filename or '')[1].lower()
    extractor = _EXTRACTORS.get(ext)
    if not extractor:
        raise ExtractError(
            f"Unsupported file type '{ext or filename}'. Supported: PDF, DOCX, CSV, TXT."
        )
    text = extractor(file_bytes)
    truncated = len(text) > MAX_CHARS
    if truncated:
        text = text[:MAX_CHARS]
    return text, truncated
